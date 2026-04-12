"""Generate QA Verification Report as .docx"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Arial'
    h.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

# Title
title = doc.add_heading('CivicRecords AI', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('QA Verification & Testing Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.style.font.size = Pt(14)

# Meta
meta = doc.add_table(rows=4, cols=2)
meta.style = 'Light Shading Accent 1'
for i, (k, v) in enumerate([
    ('Project', 'CivicRecords AI v0.1.0'),
    ('Date', 'April 12, 2026'),
    ('Prepared by', 'Claude Opus 4.6 (automated QA)'),
    ('Environment', 'Windows 11 Pro, Docker Desktop 4.68.0, Python 3.12.8'),
]):
    meta.rows[i].cells[0].text = k
    meta.rows[i].cells[1].text = v

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'CivicRecords AI has completed its initial development cycle covering all 5 sub-projects. '
    'This report documents the overnight QA sprint conducted per the project\'s development standards. '
    'Eight sprints were completed covering UI fixes, UX polish, real-data testing, end-to-end browser '
    'validation, security review, repo hygiene, documentation, and final verification.'
)
p = doc.add_paragraph()
r = p.add_run('Final status: ')
r.bold = True
p.add_run('80/80 automated tests passing. All 8 frontend pages functional. Search, ingestion, '
          'exemption detection, and request workflow verified with real municipal test data. '
          'Zero application JavaScript errors in browser console. Ready for human review before GitHub push.')

# Project Statistics
doc.add_heading('Project Statistics', level=1)
stats = doc.add_table(rows=11, cols=2)
stats.style = 'Light Shading Accent 1'
stats.alignment = WD_TABLE_ALIGNMENT.CENTER
stats_data = [
    ('Metric', 'Value'),
    ('Total commits', '59'),
    ('Total files', '150'),
    ('Total lines of code', '16,258'),
    ('Python backend files', '~70'),
    ('React frontend pages', '8'),
    ('Database tables', '10'),
    ('Database migrations', '6'),
    ('Docker services', '7'),
    ('Automated tests', '80'),
    ('Test pass rate', '100%'),
]
for i, (k, v) in enumerate(stats_data):
    stats.rows[i].cells[0].text = k
    stats.rows[i].cells[1].text = v
    if i == 0:
        for cell in stats.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

# Sprint sections
sprints = [
    ('Sprint 1: Fix Broken UI', [
        'Issues Found:',
        '- Users page returned "Not Found" \u2014 GET /users endpoint doesn\'t exist in fastapi-users v15',
        '- SPA routing worked (nginx try_files correct) but the API endpoint was missing',
        '',
        'Fixes Applied:',
        '- Created GET /admin/users endpoint in admin router with admin-only authorization',
        '- Updated Users.tsx to call /admin/users instead of /users',
        '',
        'Verification: All 7 routes return HTTP 200 with React app loading. Users page shows user list correctly.',
    ]),
    ('Sprint 2: UI/UX Polish', [
        'Issues Found (17 total):',
        '- App.tsx used <a href> tags instead of react-router <Link> \u2014 caused full page reloads',
        '- No loading states on 6 of 8 pages',
        '- No proper empty states on 4 pages',
        '- No aria-label attributes on any tables',
        '- No responsive breakpoints on 2 pages',
        '- No active nav link highlighting',
        '',
        'Fixes Applied:',
        '- Replaced all 7 <a> tags with react-router <Link> components',
        '- Added NavLink component with active state highlighting',
        '- Added loading spinners to all data-fetching pages',
        '- Added empty state messages with helpful guidance',
        '- Added aria-label to all tables and role="alert" to all error displays',
        '- Added responsive grid breakpoints throughout',
        '- Created globals.css with reusable component classes',
        '',
        'Verification: TypeScript zero errors. Frontend builds clean. Navigation instant. Console clean.',
    ]),
    ('Sprint 3: Real Data Testing', [
        'Issues Found:',
        '1. Celery worker tasks failed \u2014 module-level async_session_maker incompatible with forked workers',
        '2. Uploaded files not accessible by worker \u2014 no shared filesystem between containers',
        '3. _uploads data source not committed before dispatching Celery task \u2014 FK violation',
        '',
        'Fixes Applied:',
        '1. Created fresh SQLAlchemy engine per task execution',
        '2. Added shared Docker volumes (uploads, cache) between api and worker',
        '3. Changed flush() to commit() + refresh() before dispatching task',
        '',
        'Test Data: 3 realistic municipal documents (water quality report, police incidents, council minutes)',
        '',
        'Results:',
        '- 3 documents ingested (4 chunks, all completed)',
        '- Search returns correctly ranked results',
        '- Exemption scanner detected 9 flags: SSN, emails, phone numbers, DOB, CORA keywords',
    ]),
    ('Sprint 4: End-to-End Browser Testing', [
        'Full clerk workflow verified in browser:',
        '1. Login \u2192 Dashboard shows all services green',
        '2. Search returns ranked results with keyword highlighting',
        '3. Requests page shows status badges and deadline tracking',
        '4. Request Detail shows attached documents and workflow buttons',
        '5. Exemptions dashboard shows 9 flags and acceptance rate',
        '6. Ingestion dashboard shows 3 documents, 4 chunks',
        '7. Users page shows role badges',
        '8. Navigation between all pages is instant (SPA routing)',
        '',
        'Browser Console: Zero application errors (8 Chrome extension messages only).',
    ]),
    ('Sprint 5: Security Review', [
        'Scan Results:',
        '- .env in .gitignore, not tracked by git',
        '- No dangerouslySetInnerHTML (no XSS vectors)',
        '- No hardcoded API keys or credentials',
        '- API keys hashed (SHA-256) before storage',
        '- No cloud API keys in codebase',
        '',
        'Issue Found & Fixed:',
        '- reset_password_token_secret was hardcoded \u2192 changed to use env JWT_SECRET',
    ]),
    ('Sprint 6: Repo Hygiene', [
        'Files Created:',
        '- LICENSE \u2014 Apache License 2.0',
        '- CHANGELOG.md \u2014 Keep a Changelog format, v0.1.0',
        '- CONTRIBUTING.md \u2014 Dev setup, coding standards, PR process',
        '- README.md \u2014 Architecture diagram, quick start, configuration reference',
    ]),
    ('Sprint 7: Documentation Artifacts', [
        'Files Created:',
        '- USER-MANUAL.md \u2014 576 lines, 10 sections for non-technical staff',
        '- docs/index.html \u2014 674 lines, responsive GitHub Pages landing page',
    ]),
    ('Sprint 8: Final Verification', [
        'Test Suite: 80/80 PASSED (25.75 seconds)',
        '',
        'Test breakdown by file:',
        '- test_parsers.py: 7 tests',
        '- test_chunker.py: 7 tests',
        '- test_embedder.py: 4 tests',
        '- test_search_engine.py: 6 tests',
        '- test_search_api.py: 5 tests',
        '- test_health.py: 1 test',
        '- test_auth.py: 5 tests',
        '- test_audit.py: 5 tests',
        '- test_admin.py: 3 tests',
        '- test_service_accounts.py: 4 tests',
        '- test_datasources.py: 4 tests',
        '- test_documents.py: 2 tests',
        '- test_pipeline.py: 2 tests',
        '- test_requests.py: 9 tests',
        '- test_exemptions.py: 16 tests',
        '',
        'Frontend: TypeScript zero errors. Build: 225.68 KB JS, 17.35 KB CSS.',
        'Docker: All 7 services running and healthy.',
    ]),
]

for title, lines in sprints:
    doc.add_heading(title, level=1)
    for line in lines:
        if line == '':
            doc.add_paragraph()
        elif line.startswith('- ') or line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. ') or line.startswith('6. ') or line.startswith('7. ') or line.startswith('8. '):
            doc.add_paragraph(line[2:] if line.startswith('- ') else line, style='List Bullet' if line.startswith('- ') else 'List Number')
        else:
            p = doc.add_paragraph()
            if line.endswith(':'):
                r = p.add_run(line)
                r.bold = True
            else:
                p.add_run(line)

# Known Limitations
doc.add_heading('Known Limitations', level=1)
limitations = [
    'Search score display \u2014 RRF fusion scores show as 1.1-1.6%. Could be normalized to 0-100%.',
    'File names in ingestion \u2014 Uploaded files get UUID prefix. Original name should be preserved.',
    'shadcn/ui not implemented \u2014 Current UI uses raw Tailwind. Functional but less polished.',
    'No file upload UI \u2014 Documents uploaded via API only. Drag-and-drop widget needed.',
    'Audit retention enforcement \u2014 Config exists but no cleanup job implemented.',
    'Docker postgres port \u2014 Integration tests need dev compose override to expose port.',
    'JWT key length warning \u2014 Test default is short. Production uses proper 32-byte key.',
]
for i, lim in enumerate(limitations, 1):
    doc.add_paragraph(f'{i}. {lim}', style='List Number')

# Recommendations
doc.add_heading('Recommendations for Next Steps', level=1)
recs = [
    'Normalize search scores for display (map RRF to 0-100% range)',
    'Add drag-and-drop file upload to Sources page',
    'Install shadcn/ui for component polish',
    'Add audit retention cleanup Celery task',
    'Generate README-FULL.pdf with architecture diagrams',
    'Push to GitHub after human review',
    'Enable GitHub Discussions and seed with welcome posts',
    'Add state exemption rules for 4 more pilot states beyond Colorado',
]
for r in recs:
    doc.add_paragraph(r, style='List Bullet')

# Save
doc.save('C:/Users/scott/Desktop/Claude/civicrecords-ai/QA-VERIFICATION-REPORT.docx')
print('Report saved to QA-VERIFICATION-REPORT.docx')

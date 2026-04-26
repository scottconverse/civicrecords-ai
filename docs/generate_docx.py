"""
CivicRecords AI — DOCX Generator
Produces README.docx and USER-MANUAL.docx at the repository root.
Run: python docs/generate_docx.py
"""

from pathlib import Path
from datetime import date
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
REPO_ROOT    = Path(__file__).parent.parent
DOCS_DIR     = Path(__file__).parent
USER_MANUAL  = REPO_ROOT / "USER-MANUAL.md"
README_MD    = REPO_ROOT / "README.md"

# Markdown image-ref pattern: ![alt](path)
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def _resolve_png(md_dir: Path, ref_path: str) -> Path:
    """Resolve a markdown image ref to a PNG on disk.

    python-docx cannot embed SVG, so prefer the .png sibling when the ref is .svg.
    Path is resolved relative to the .md file's directory.
    """
    p = (md_dir / ref_path).resolve()
    if p.suffix.lower() == ".svg":
        p = p.with_suffix(".png")
    return p


def embed_markdown_image(doc, md_dir: Path, alt: str, ref_path: str, width_inches: float = 6.0):
    """Embed an image referenced by markdown ![alt](path) into a python-docx document.

    Skips with a logged warning if the resolved PNG is missing (does not crash).
    Adds an italic centered caption below the image when alt text is present.
    """
    png = _resolve_png(md_dir, ref_path)
    if not png.exists():
        print(f"  WARNING: image not found, skipping: {png} (ref: {ref_path})")
        return
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(png), width=Inches(width_inches))
        if alt:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            crun = cap.add_run(alt)
            crun.font.name = "Calibri"
            crun.font.size = Pt(9)
            crun.font.italic = True
            crun.font.color.rgb = GRAY_TEXT
    except Exception as e:
        print(f"  WARNING: failed to embed image {png}: {e}")

# ---------------------------------------------------------------------------
# Brand colours  (RGB tuples)
# ---------------------------------------------------------------------------
CIVIC_DARK   = RGBColor(0x1a, 0x3a, 0x5c)
CIVIC_MID    = RGBColor(0x25, 0x63, 0xeb)
CIVIC_ACCENT = RGBColor(0x3b, 0x82, 0xf6)
CIVIC_LIGHT  = RGBColor(0xef, 0xf6, 0xff)
CIVIC_ROW    = RGBColor(0xf8, 0xfa, 0xfc)
WHITE        = RGBColor(0xff, 0xff, 0xff)
GRAY_BG      = RGBColor(0xf3, 0xf4, 0xf6)
GRAY_TEXT    = RGBColor(0x6b, 0x72, 0x80)

# Hex strings for XML-level shading
HEX_DARK     = "1a3a5c"
HEX_MID      = "2563eb"
HEX_LIGHT    = "eff6ff"
HEX_ROW      = "f8fafc"
HEX_GRAY     = "f3f4f6"
HEX_WHITE    = "FFFFFF"

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, sides=("top", "bottom", "left", "right"), color="d1d5db", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    str(sz))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_toc_field(doc):
    """Insert a Word TOC field that Word will populate on open/update."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run = para.add_run()
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)
    return para


def add_page_number(footer_para):
    """Insert {PAGE} field in a paragraph."""
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run = footer_para.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def add_section_break(doc):
    """Insert a continuous section break."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    type_elem = OxmlElement("w:type")
    type_elem.set(qn("w:val"), "nextPage")
    sectPr.append(type_elem)
    pPr.append(sectPr)
    return para


# ---------------------------------------------------------------------------
# Document setup helpers
# ---------------------------------------------------------------------------
def setup_document(title_text, subtitle_text=None):
    """Create a Document with standard margins, header, and footer."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(0.85)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    return doc


def add_header_footer(section, header_text, include_page_num=True):
    """Apply header and footer to a document section."""
    section.different_first_page_header_footer = False

    # Header
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        hp = header.paragraphs[0]
    else:
        hp = header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(header_text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(9)
    run.font.color.rgb = CIVIC_MID
    run.font.italic = True

    # Thin blue rule under header via bottom border on header paragraph
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), HEX_MID)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        fp = footer.paragraphs[0]
    else:
        fp = footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Apache License 2.0  —  github.com/CivicSuite/civicrecords-ai").font.size = Pt(8)
    if include_page_num:
        fp.add_run("     Page ").font.size = Pt(8)
        add_page_number(fp)


def apply_h1_style(para):
    para.style = "Heading 1"
    run = para.runs[0] if para.runs else para.add_run()
    run.font.name  = "Calibri"
    run.font.size  = Pt(16)
    run.font.bold  = True
    run.font.color.rgb = CIVIC_DARK


def apply_h2_style(para):
    para.style = "Heading 2"
    run = para.runs[0] if para.runs else para.add_run()
    run.font.name  = "Calibri"
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = CIVIC_MID


def apply_h3_style(para):
    para.style = "Heading 3"
    run = para.runs[0] if para.runs else para.add_run()
    run.font.name  = "Calibri"
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = CIVIC_DARK


def body_para(doc, text, indent=False, italic=False, bold=False, color=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.bold   = bold
    if color:
        run.font.color.rgb = color
    if indent:
        para.paragraph_format.left_indent = Inches(0.25)
    return para


def bullet_para(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    para.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    para.paragraph_format.space_after  = Pt(2)
    return para


def numbered_para(doc, text, level=0):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    para.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    para.paragraph_format.space_after = Pt(2)
    return para


def code_para(doc, text):
    """Courier New 9pt, light gray background."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(1)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.left_indent  = Inches(0.25)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    # Gray background via paragraph shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), HEX_GRAY)
    pPr.append(shd)
    return para


def mermaid_placeholder(doc, diagram_title):
    """Styled blue-bordered paragraph replacing a mermaid code block."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.left_indent  = Inches(0.5)
    para.paragraph_format.right_indent = Inches(0.5)

    run_title = para.add_run(f"Architecture Diagram: {diagram_title}")
    run_title.font.name   = "Calibri"
    run_title.font.size   = Pt(11)
    run_title.font.bold   = True
    run_title.font.color.rgb = CIVIC_MID

    para.add_run("\n")

    run_note = para.add_run("(See docs/diagrams/ for Mermaid source — renders on GitHub)")
    run_note.font.name   = "Calibri"
    run_note.font.size   = Pt(9)
    run_note.font.italic = True
    run_note.font.color.rgb = GRAY_TEXT

    # Blue left border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "18")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), HEX_MID)
    pBdr.append(left)
    pPr.append(pBdr)

    # Light blue background
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), HEX_LIGHT)
    pPr.append(shd)

    return para


def make_table(doc, headers, rows, col_widths=None):
    """Create a styled table with civic blue header row and alternating row shading."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, HEX_DARK)
        set_cell_borders(cell, color=HEX_DARK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(text)
        run.font.name  = "Calibri"
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = HEX_WHITE if r_idx % 2 == 0 else HEX_ROW
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(cell_text))
            run.font.name = "Calibri"
            run.font.size = Pt(9)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    return table


# ---------------------------------------------------------------------------
# Data constants (shared with PDF generator)
# ---------------------------------------------------------------------------
FEATURES = [
    ("AI-Powered Search",
     "Natural language hybrid search (semantic + keyword) across all ingested documents with "
     "source attribution, normalized relevance scores, and optional AI-generated summaries."),
    ("Document Ingestion",
     "Automatic parsing of PDF, DOCX, XLSX, CSV, email, HTML, and text files. Scanned "
     "documents processed via multimodal AI (Gemma 4) with Tesseract OCR fallback."),
    ("Exemption Detection",
     "Tier 1 PII detection (SSN, credit card with Luhn validation, phone, email, bank "
     "accounts, state-specific driver's licenses) plus per-state statutory keyword matching. "
     "All flags require human confirmation."),
    ("Request Management",
     "Full lifecycle tracking with 11 statuses. Timeline, messaging, fee tracking, and "
     "response letter generation."),
    ("Guided Onboarding",
     "3-phase wizard: city profile, systems identification across 12 municipal domains, gap map."),
    ("Municipal Systems Catalog",
     "25+ municipal software vendors across 12 functional domains with discovery hints."),
    ("Universal Connector Framework",
     "Standardized 4-method protocol. Ships with file system, REST API, and ODBC connectors."),
    ("Scheduled Sync & Idempotent Ingestion",
     "Per-source cron scheduling with 5-minute floor. Idempotent dedup by content hash or "
     "stable source-path. Concurrent-update races prevented via SELECT FOR UPDATE."),
    ("Sync Failure Tracking & Circuit Breaker",
     "Two-layer retry model. Automatic circuit breaker after 5 consecutive full-run failures. "
     "Admin UI with colored health badge, failed records panel, and Sync Now button."),
    ("Operational Analytics",
     "Real-time metrics: average response time, deadline compliance rate, overdue requests."),
    ("Notification Service",
     "12 template-based email notifications with SMTP delivery via Celery beat."),
    ("Compliance by Design",
     "Hash-chained audit logs, human-in-the-loop enforcement, AI content labeling, "
     "data sovereignty verification. 50-state regulatory compliance design."),
    ("Civic Design System",
     "shadcn/ui, civic blue design tokens, WCAG 2.2 AA targeted."),
    ("Federation-Ready",
     "REST API with service accounts enables cross-jurisdiction record discovery."),
    ("50-State Exemption Rules",
     "180 rules across 51 jurisdictions with per-state statutory keyword matching."),
    ("Department Access Controls",
     "Staff scoped to own department; admins see all. Full RBAC role hierarchy."),
    ("Compliance Templates",
     "5 compliance documents with city profile variable substitution."),
]

TECH_STACK = [
    ["Layer",       "Technology",               "Version"],
    ["Language",    "Python",                   "3.12"],
    ["API",         "FastAPI",                  "0.115+"],
    ["ORM",         "SQLAlchemy",               "2.0"],
    ["Database",    "PostgreSQL + pgvector",    "17"],
    ["Queue",       "Redis",                    "7.2"],
    ["Task runner", "Celery",                   "5.x"],
    ["LLM runtime", "Ollama",                   "latest"],
    ["Frontend",    "React + shadcn/ui",        "18"],
    ["Containers",  "Docker Compose",           "v2"],
    ["Testing",     "pytest / vitest",          "latest"],
]


# ---------------------------------------------------------------------------
# README.docx
# ---------------------------------------------------------------------------
def build_readme_docx(out_path):
    print(f"  Generating {out_path.name}...")
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    add_header_footer(
        doc.sections[0],
        "CivicRecords AI  |  v1.1+  |  Open-Source Municipal Records AI"
    )

    # Title block
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("CivicRecords AI")
    title_run.font.name  = "Calibri"
    title_run.font.size  = Pt(36)
    title_run.font.bold  = True
    title_run.font.color.rgb = CIVIC_DARK

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Open-Source, Locally-Hosted AI for Municipal Open Records")
    sub_run.font.name  = "Calibri"
    sub_run.font.size  = Pt(14)
    sub_run.font.color.rgb = CIVIC_MID

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(f"v1.1+  ·  {date.today().strftime('%B %Y')}  ·  Apache 2.0")
    meta_run.font.name   = "Calibri"
    meta_run.font.size   = Pt(11)
    meta_run.font.italic = True
    meta_run.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()  # spacer

    # --- Why This Exists ---
    h = doc.add_heading("Why This Exists", level=1)
    apply_h1_style(h)
    body_para(doc,
        "Every city in America processes open records requests (FOIA, CORA, and state equivalents). "
        "Staff manually search file shares, email archives, and databases — then review every document "
        "for exemptions before release. It is slow, error-prone, and a growing burden as request "
        "volumes increase. No open-source tool existed for the responder side of open records at the "
        "municipal level. CivicRecords AI fills that gap."
    )

    # --- Key Features ---
    h = doc.add_heading("Key Features", level=1)
    apply_h1_style(h)
    body_para(doc, f"CivicRecords AI ships with {len(FEATURES)} documented features:")
    doc.add_paragraph()

    feat_rows = [[f["name"], f["desc"]] for f in [{"name": n, "desc": d} for n, d in FEATURES]]
    make_table(
        doc,
        ["Feature", "Description"],
        feat_rows,
        col_widths=[Inches(1.7), Inches(4.8)],
    )
    doc.add_paragraph()

    # --- Quick Start ---
    h = doc.add_heading("Quick Start", level=1)
    apply_h1_style(h)

    h2 = doc.add_heading("Requirements", level=2)
    apply_h2_style(h2)
    for item in [
        "Docker Desktop (Windows 10/11, macOS 13+) or Docker Engine (Linux)",
        "8+ CPU cores, 32 GB RAM, 50 GB free disk space",
        "No internet connection required after initial setup",
    ]:
        bullet_para(doc, item)

    h2 = doc.add_heading("Install (Windows)", level=2)
    apply_h2_style(h2)
    for line in [
        "git clone https://github.com/CivicSuite/civicrecords-ai.git",
        "cd civicrecords-ai",
        ".\\install.ps1",
    ]:
        code_para(doc, line)

    h2 = doc.add_heading("Install (macOS / Linux)", level=2)
    apply_h2_style(h2)
    for line in [
        "git clone https://github.com/CivicSuite/civicrecords-ai.git",
        "cd civicrecords-ai",
        "bash install.sh",
    ]:
        code_para(doc, line)

    h2 = doc.add_heading("First Use", level=2)
    apply_h2_style(h2)
    for step in [
        "Open http://localhost:8080 in your browser",
        "Sign in with the admin credentials configured in .env",
        "Go to Sources → Add Source → enter a directory path to your documents",
        "Click Ingest Now — documents are parsed, chunked, and indexed automatically",
        "Go to Search — type a natural language query and get cited results",
    ]:
        numbered_para(doc, step)

    # --- Architecture ---
    h = doc.add_heading("Architecture", level=1)
    apply_h1_style(h)
    body_para(doc,
        "CivicRecords AI deploys as a 7-service Docker Compose stack. All services run in "
        "Linux containers regardless of host operating system."
    )
    doc.add_paragraph()

    # Architecture as a formatted table (no image — SVG not embeddable without Pillow)
    make_table(
        doc,
        ["Layer", "Service", "Technology", "Port"],
        [
            ["Presentation", "frontend",  "nginx + React 18 + shadcn/ui",    "8080"],
            ["API",          "api",       "Python 3.12 / FastAPI",           "8000"],
            ["Workers",      "worker",    "Celery worker (ingestion tasks)", "—"],
            ["Scheduler",    "beat",      "Celery beat (cron triggers)",     "—"],
            ["Data store",   "postgres",  "PostgreSQL 17 + pgvector",        "5432"],
            ["Queue",        "redis",     "Redis 7.2",                       "6379"],
            ["Inference",    "ollama",    "Ollama local LLM runtime",        "11434"],
        ],
        col_widths=[Inches(1.1), Inches(1.1), Inches(2.7), Inches(0.7)],
    )
    doc.add_paragraph()
    body_para(doc,
        "Tech stack: Python 3.12, FastAPI, SQLAlchemy 2.0, React 18, shadcn/ui, Tailwind CSS, "
        "Alembic, Celery, pgvector, nomic-embed-text, Gemma 4.",
        italic=True
    )

    # Embed architecture diagrams referenced in README.md
    for alt, ref in [
        ("Deployment stack — entire system runs inside Docker Compose on the city's network. "
         "No cloud, no outbound by default.",
         "docs/diagrams/deployment-stack.svg"),
        ("LLM call flow: records-ai routes through civiccore.llm to a local Ollama provider; "
         "cloud providers are opt-in only.",
         "docs/diagrams/llm-flow.svg"),
        ("Sovereignty boundary: all runtime components live inside the city's on-prem network. "
         "Connector credentials encrypted at rest with Fernet.",
         "docs/diagrams/sovereignty.svg"),
    ]:
        embed_markdown_image(doc, REPO_ROOT, alt, ref)

    # --- Technology Stack ---
    h = doc.add_heading("Technology Stack", level=1)
    apply_h1_style(h)
    make_table(
        doc,
        ["Layer", "Technology", "Version"],
        [row for row in TECH_STACK[1:]],  # skip header row
        col_widths=[Inches(1.3), Inches(2.4), Inches(1.0)],
    )
    doc.add_paragraph()

    # --- Configuration ---
    h = doc.add_heading("Configuration", level=1)
    apply_h1_style(h)
    body_para(doc, "All configuration is via environment variables in .env:")
    doc.add_paragraph()
    make_table(
        doc,
        ["Variable", "Description", "Default"],
        [
            ["DATABASE_URL",         "PostgreSQL connection string",     "postgresql+asyncpg://..."],
            ["JWT_SECRET",           "Secret key for JWT tokens",        "(must be set)"],
            ["FIRST_ADMIN_EMAIL",    "Initial admin account email",      "admin@example.gov"],
            ["FIRST_ADMIN_PASSWORD", "Initial admin account password",   "(must be set)"],
            ["OLLAMA_BASE_URL",      "Ollama API endpoint",              "http://ollama:11434"],
            ["REDIS_URL",            "Redis connection string",          "redis://redis:6379/0"],
            ["AUDIT_RETENTION_DAYS","Audit log retention period",       "1095 (3 years)"],
        ],
        col_widths=[Inches(2.0), Inches(2.5), Inches(2.1)],
    )

    # --- User Roles ---
    h = doc.add_heading("User Roles", level=1)
    apply_h1_style(h)
    make_table(
        doc,
        ["Role", "Permissions", "Status"],
        [
            ["Admin",     "Full access: users, system config, rules, audit logs, onboarding", "Built"],
            ["Staff",     "Search, create requests, attach documents, review flags, fees",    "Built"],
            ["Reviewer",  "Everything Staff + approve/reject responses and flags",            "Built"],
            ["Read-Only", "View search results and request status only",                      "Built"],
            ["Liaison",   "Scoped to assigned department, attach documents and notes",        "MVP-NOW"],
            ["Public",    "Submit requests, track own requests, search published records",    "v1.1"],
        ],
        col_widths=[Inches(1.0), Inches(4.0), Inches(1.0)],
    )

    # --- Status ---
    h = doc.add_heading("Status", level=1)
    apply_h1_style(h)
    body_para(doc,
        "v1.3.0 — Phase 1 CivicCore extraction landed. civiccore 0.1.0 is now a "
        "release-wheel dependency; two-layer migration order (civiccore first via "
        "subprocess, then records). Infrastructure-only release — no API or UI changes."
    )
    body_para(doc,
        "v1.2.0 — Tier 5 + Tier 6 complete. Adds minimal public request portal (T5D), "
        "unsigned Windows installer (T5E), and ENG-001 closed "
        "(Tier 6 at-rest Fernet encryption on data_sources.connection_config shipped 2026-04-23). "
        "Builds on Phase 2 department access controls, 50-state exemption rules, and compliance templates."
    )
    for item in [
        "29 database tables managed by 16 Alembic migration scripts",
        "~30 REST API endpoints under /api/v1/",
        "617 backend + 36 frontend automated tests passing",
        "Tested on Windows 11 (Docker Desktop) and Ubuntu 22.04 (Docker Engine)",
        "AMD GPU/NPU hardware auto-detection (ROCm on Linux, DirectML on Windows)",
    ]:
        bullet_para(doc, item)

    # --- License ---
    h = doc.add_heading("License", level=1)
    apply_h1_style(h)
    body_para(doc,
        "Apache License 2.0 — see LICENSE. All dependencies use permissive (MIT, Apache 2.0, BSD) "
        "or weak-copyleft (LGPL, MPL) licenses. No AGPL, SSPL, or BSL dependencies."
    )
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("github.com/CivicSuite/civicrecords-ai")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = CIVIC_MID

    doc.save(str(out_path))
    print(f"  Done: {out_path}")


# ---------------------------------------------------------------------------
# USER-MANUAL.docx  (parsed from USER-MANUAL.md)
# ---------------------------------------------------------------------------

def extract_mermaid_title(block_lines):
    """Try to pull a title from the first comment line of a mermaid block."""
    for line in block_lines:
        stripped = line.strip()
        if stripped.startswith("%%") or stripped.startswith("title"):
            return stripped.lstrip("%").strip().lstrip("title").strip()
        if stripped and not stripped.startswith("graph") and not stripped.startswith("sequenceDiagram") \
                and not stripped.startswith("classDiagram") and not stripped.startswith("flowchart"):
            return stripped[:60]
    return "System Diagram"


def parse_inline_formatting(text):
    """Return list of (run_text, bold, italic, code) tuples for inline MD formatting."""
    # Simple parser: handle **bold**, *italic*, `code`
    segments = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            segments.append((text[last:m.start()], False, False, False))
        if m.group().startswith("**"):
            segments.append((m.group(2), True, False, False))
        elif m.group().startswith("*"):
            segments.append((m.group(3), False, True, False))
        elif m.group().startswith("`"):
            segments.append((m.group(4), False, False, True))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False, False, False))
    return segments


def add_formatted_para(doc, text, style_name="Normal", indent=False, space_after=4):
    """Add a paragraph with inline bold/italic/code formatting from markdown."""
    para = doc.add_paragraph(style=style_name)
    para.paragraph_format.space_after = Pt(space_after)
    if indent:
        para.paragraph_format.left_indent = Inches(0.25)

    segments = parse_inline_formatting(text)
    for (seg_text, bold, italic, code) in segments:
        run = para.add_run(seg_text)
        run.font.name  = "Courier New" if code else "Calibri"
        run.font.size  = Pt(9) if code else Pt(11)
        run.font.bold  = bold
        run.font.italic = italic
        if code:
            run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    return para


def parse_table_row(line):
    """Split a markdown table row into cells."""
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def build_user_manual_docx(out_path):
    if not USER_MANUAL.exists():
        print(f"  WARNING: {USER_MANUAL} not found — skipping USER-MANUAL.docx")
        return

    print(f"  Generating {out_path.name}...")
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    add_header_footer(
        doc.sections[0],
        "CivicRecords AI User Manual  |  v1.1+"
    )

    lines = USER_MANUAL.read_text(encoding="utf-8").splitlines()

    i = 0
    in_code_block  = False
    code_lang      = ""
    code_lines     = []
    in_table       = False
    table_headers  = []
    table_rows     = []
    first_h1_seen  = False

    # Inject title page before parsing markdown
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("CivicRecords AI")
    title_run.font.name  = "Calibri"
    title_run.font.size  = Pt(36)
    title_run.font.bold  = True
    title_run.font.color.rgb = CIVIC_DARK

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("User Manual")
    sub_run.font.name  = "Calibri"
    sub_run.font.size  = Pt(20)
    sub_run.font.color.rgb = CIVIC_MID

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(f"v1.1+  ·  For Municipal Records Staff  ·  {date.today().strftime('%B %Y')}")
    meta_run.font.name   = "Calibri"
    meta_run.font.size   = Pt(11)
    meta_run.font.italic = True
    meta_run.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()

    # TOC placeholder
    toc_h = doc.add_heading("Table of Contents", level=1)
    apply_h1_style(toc_h)
    add_toc_field(doc)
    doc.add_page_break()

    def flush_table():
        nonlocal in_table, table_headers, table_rows
        if table_headers and table_rows:
            make_table(doc, table_headers, table_rows)
            doc.add_paragraph()
        in_table = False
        table_headers = []
        table_rows    = []

    def flush_code(lang, lines_list):
        for cl in lines_list:
            code_para(doc, cl)
        doc.add_paragraph()

    while i < len(lines):
        line = lines[i]

        # --- Code block toggle ---
        if line.startswith("```"):
            if in_code_block:
                if code_lang == "mermaid":
                    title = extract_mermaid_title(code_lines)
                    mermaid_placeholder(doc, title)
                else:
                    flush_code(code_lang, code_lines)
                in_code_block = False
                code_lang  = ""
                code_lines = []
            else:
                if in_table:
                    flush_table()
                in_code_block = True
                code_lang  = line[3:].strip().lower()
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Markdown table ---
        if line.startswith("|"):
            if not in_table:
                in_table = True
                table_headers = parse_table_row(line)
                i += 1
                # Skip separator row (---|---|---)
                if i < len(lines) and re.match(r"^\|[-| :]+\|", lines[i]):
                    i += 1
                continue
            else:
                if re.match(r"^\|[-| :]+\|", line):
                    i += 1
                    continue
                table_rows.append(parse_table_row(line))
                i += 1
                continue
        else:
            if in_table:
                flush_table()

        # --- Blank line ---
        if not line.strip():
            i += 1
            continue

        # --- Markdown image: ![alt](path) ---
        img_m = IMAGE_RE.match(line)
        if img_m:
            embed_markdown_image(doc, USER_MANUAL.parent, img_m.group(1), img_m.group(2))
            i += 1
            continue

        # --- Horizontal rule ---
        if line.strip() in ("---", "***", "___"):
            # Insert a thin blue paragraph border as separator
            sep_para = doc.add_paragraph()
            pPr = sep_para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), HEX_MID)
            pBdr.append(bottom)
            pPr.append(pBdr)
            sep_para.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # --- Headings ---
        if line.startswith("# "):
            text = line[2:].strip()
            # Skip the original "# CivicRecords AI — User Manual" (title already added)
            if not first_h1_seen:
                first_h1_seen = True
                # Still add as H1 for TOC
                h = doc.add_heading(text, level=1)
                apply_h1_style(h)
            else:
                h = doc.add_heading(text, level=1)
                apply_h1_style(h)
            i += 1
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            h = doc.add_heading(text, level=2)
            apply_h2_style(h)
            i += 1
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            h = doc.add_heading(text, level=3)
            apply_h3_style(h)
            i += 1
            continue

        if line.startswith("#### "):
            text = line[5:].strip()
            h = doc.add_heading(text, level=3)
            apply_h3_style(h)
            i += 1
            continue

        # --- Blockquote ---
        if line.startswith("> "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent  = Inches(0.3)
            para.paragraph_format.right_indent = Inches(0.3)
            para.paragraph_format.space_after  = Pt(4)
            run = para.add_run(text)
            run.font.name   = "Calibri"
            run.font.size   = Pt(11)
            run.font.italic = True
            run.font.color.rgb = CIVIC_MID
            # Left blue border
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"),   "single")
            left.set(qn("w:sz"),    "12")
            left.set(qn("w:space"), "6")
            left.set(qn("w:color"), HEX_MID)
            pBdr.append(left)
            pPr.append(pBdr)
            i += 1
            continue

        # --- Bullet list (- or * or numbered) ---
        bullet_m = re.match(r"^(\s*)[-*+] (.+)$", line)
        if bullet_m:
            indent_level = len(bullet_m.group(1)) // 2
            text = bullet_m.group(2).strip()
            bullet_para(doc, text, level=indent_level)
            i += 1
            continue

        numbered_m = re.match(r"^(\s*)\d+\. (.+)$", line)
        if numbered_m:
            indent_level = len(numbered_m.group(1)) // 2
            text = numbered_m.group(2).strip()
            numbered_para(doc, text, level=indent_level)
            i += 1
            continue

        # --- Regular paragraph ---
        # Collect continuation lines
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if (not next_line.strip()
                    or next_line.startswith("#")
                    or next_line.startswith(">")
                    or next_line.startswith("```")
                    or next_line.startswith("|")
                    or re.match(r"^(\s*)[-*+] ", next_line)
                    or re.match(r"^(\s*)\d+\. ", next_line)
                    or next_line.strip() in ("---", "***", "___")):
                break
            para_lines.append(next_line)
            i += 1

        full_text = " ".join(para_lines)
        add_formatted_para(doc, full_text)

    # Flush any remaining table
    if in_table:
        flush_table()

    doc.save(str(out_path))
    print(f"  Done: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    print("CivicRecords AI DOCX Generator")
    print("=" * 40)

    readme_docx     = REPO_ROOT / "README.docx"
    usermanual_docx = REPO_ROOT / "USER-MANUAL.docx"

    build_readme_docx(readme_docx)
    build_user_manual_docx(usermanual_docx)

    print()
    print("Generated:")
    for p in [readme_docx, usermanual_docx]:
        if p.exists():
            size_kb = p.stat().st_size // 1024
            print(f"  {p}  ({size_kb} KB)")
        else:
            print(f"  {p}  (NOT FOUND — check errors above)")


if __name__ == "__main__":
    main()

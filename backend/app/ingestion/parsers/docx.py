import logging
import tempfile
import zipfile
from pathlib import Path

from docx import Document as DocxDocument

from app.ingestion.parsers.base import BaseParser, ParsedPage, ParseResult

logger = logging.getLogger(__name__)

# ZIP entries that indicate embedded VBA macros in DOCX files.
_DOCX_VBA_ENTRIES = {"word/vbaproject.bin", "word/vbadata.xml"}


class DocxParser(BaseParser):
    supported_extensions = [".docx", ".docm"]

    @staticmethod
    def _strip_macros(file_path: Path) -> tuple[Path | None, list[str]]:
        """Check for VBA macro entries and return a sanitized copy if found.

        Returns (sanitized_path, stripped_entry_names).
        sanitized_path is None if no macros were found.
        """
        try:
            with zipfile.ZipFile(file_path, "r") as zin:
                found = [n for n in zin.namelist() if n.lower() in _DOCX_VBA_ENTRIES]
                if not found:
                    return None, []
                tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                tmp_path = Path(tmp.name)
                tmp.close()
                with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.namelist():
                        if item.lower() not in _DOCX_VBA_ENTRIES:
                            zout.writestr(item, zin.read(item))
                logger.warning(
                    "Stripped VBA macros from %s: %s", file_path.name, found
                )
                return tmp_path, found
        except zipfile.BadZipFile:
            logger.warning("Not a valid ZIP archive: %s", file_path.name)
            return None, []

    def parse(self, file_path: Path) -> ParseResult:
        sanitized_path, stripped_entries = self._strip_macros(file_path)
        parse_path = sanitized_path or file_path
        try:
            doc = DocxDocument(str(parse_path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    paragraphs.append("\n".join(rows))
            full_text = "\n\n".join(paragraphs)
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
            if doc.core_properties:
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
                if doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
            if stripped_entries:
                metadata["macros_stripped"] = True
                metadata["stripped_entries"] = stripped_entries
            return ParseResult(
                pages=[ParsedPage(text=full_text, page_number=1)],
                metadata=metadata,
                file_type="docx",
            )
        finally:
            if sanitized_path and sanitized_path.exists():
                sanitized_path.unlink()
